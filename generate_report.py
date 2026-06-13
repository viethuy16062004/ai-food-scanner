import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="none"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="888888"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>'
    )
    tblBorders = parse_xml(borders_xml)
    tblPr.append(tblBorders)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(6, 95, 70) # Emerald 800
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(15, 118, 110) # Teal 700
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(55, 65, 81) # Gray 700
    return p

def main():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles config
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(31, 41, 55) # Gray 800
    
    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_school = p_school.add_run("TRƯỜNG ĐẠI HỌC BÁCH KHOA - ĐHQG TP.HCM\nKHOA KHOA HỌC VÀ KỸ THUẬT MÁY TÍNH\n---***---")
    run_school.font.size = Pt(12)
    run_school.font.name = 'Arial'
    run_school.bold = True
    run_school.font.color.rgb = RGBColor(55, 65, 81)
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BÁO CÁO BÀI TẬP LỚN CHI TIẾT\nMÔN: CÔNG NGHỆ PHẦN MỀM")
    run_title.font.size = Pt(18)
    run_title.font.name = 'Arial'
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(6, 95, 70)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("ĐỀ TÀI: PHÂN TÍCH VÀ XÂY DỰNG HỆ THỐNG\nAI NUTRISCAN - QUÉT VÀ THEO DÕI DINH DƯỠNG THỰC PHẨM")
    run_sub.font.size = Pt(14)
    run_sub.font.name = 'Arial'
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(15, 118, 110)
    
    for _ in range(5):
        doc.add_paragraph()
        
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.left_indent = Inches(1.5)
    
    run_info = p_info.add_run(
        "Giảng viên hướng dẫn:\tCô/Thầy...\n"
        "Nhóm thực hiện:\t\tNhóm AI Food Scanner\n"
        "Thành viên nhóm:\n"
        "  1. Trần Việt Huy\t\tMSSV: 22... (Trưởng nhóm & Lập trình chính)\n"
        "  2. ...\t\t\tMSSV: ...\n"
        "  3. ...\t\t\tMSSV: ...\n"
    )
    run_info.font.size = Pt(11)
    run_info.font.name = 'Arial'
    run_info.font.color.rgb = RGBColor(31, 41, 55)
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("TP. HỒ CHÍ MINH, THÁNG 05 NĂM 2026")
    run_date.font.size = Pt(10)
    run_date.font.name = 'Arial'
    run_date.bold = True
    run_date.font.color.rgb = RGBColor(107, 114, 128)
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # SECTION 1: GIỚI THIỆU ĐỀ TÀI
    # ----------------------------------------------------
    add_heading_styled(doc, "1. GIỚI THIỆU ĐỀ TÀI", 1)
    
    p = doc.add_paragraph()
    p.add_run("Trong xã hội hiện đại, nhu cầu theo dõi sức khỏe và kiểm soát dinh dưỡng ngày càng trở nên quan trọng. ")
    p.add_run("Tuy nhiên, việc tự ghi nhận calo và các chất dinh dưỡng thủ công từ thực phẩm tiêu thụ hàng ngày mất rất nhiều thời gian và dễ sai sót. ")
    p.add_run("Dự án ")
    p.add_run("AI NutriScan").bold = True
    p.add_run(" được phát triển nhằm giải quyết vấn đề này. Ứng dụng tích hợp công nghệ trí tuệ nhân tạo (AI) giúp nhận diện thực phẩm qua hình ảnh, phân tích lập tức hàm lượng calo, protein, carbohydrate, fat và đưa ra điểm số đánh giá độ lành mạnh. ")
    p.add_run("Từ đó, hệ thống cung cấp các công cụ theo dõi sức khỏe dài hạn như Nhật ký sức khỏe, Lập kế hoạch bữa ăn và Danh sách mua sắm, giúp người dùng duy trì lối sống lành mạnh một cách khoa học nhất.")
    
    p2 = doc.add_paragraph()
    p2.add_run("Hệ thống được phát triển dựa trên mô hình Client-Server hiện đại: ")
    p2.add_run("Frontend").bold = True
    p2.add_run(" xây dựng trên thư viện React, ")
    p2.add_run("Backend").bold = True
    p2.add_run(" sử dụng framework Spring Boot kết hợp với hệ quản trị cơ sở dữ liệu MySQL, và một ")
    p2.add_run("Node.js Gemini Gateway").bold = True
    p2.add_run(" chuyên biệt phục vụ tối ưu hóa hình ảnh và kết nối bảo mật tới mô hình ngôn ngữ lớn (LLM) của Google - Gemini API.")

    # ----------------------------------------------------
    # SECTION 2: CÁC CHỨC NĂNG CHI TIẾT
    # ----------------------------------------------------
    add_heading_styled(doc, "2. CÁC CHỨC NĂNG CHI TIẾT CỦA HỆ THỐNG", 1)
    
    p = doc.add_paragraph()
    p.add_run("Hệ thống được thiết kế hướng tới việc chia nhỏ các chức năng để đảm bảo tính rõ ràng và chi tiết, tránh các định nghĩa chung chung. Các module chức năng cụ thể bao gồm:")
    
    features = [
        ("Module Xác thực & Tài khoản:", "Đăng ký tài khoản mới, Đăng nhập hệ thống, Đăng xuất tài khoản, Xem thông tin cá nhân, Chỉnh sửa hồ sơ cá nhân và Thay đổi mật khẩu."),
        ("Module Nhận diện món ăn bằng AI:", "Tải ảnh thực phẩm từ thiết bị, Chụp ảnh thực phẩm bằng camera, Nhận diện thực phẩm qua AI, Xem kết quả phân tích dinh dưỡng (Calo, Đạm, Tinh bột, Béo, Healthy Score) và Lưu lịch sử quét dinh dưỡng."),
        ("Module Nhật ký Sức khỏe:", "Ghi nhật ký cân nặng, Xem biểu đồ cân nặng & BMI, Ghi nhận lượng nước uống (ml), Ghi nhận thời gian vận động (phút) và Xem biểu đồ dinh dưỡng tuần."),
        ("Module Kế hoạch Bữa ăn:", "Lập thực đơn bữa ăn, Xem thực đơn bữa ăn, Thêm món ăn vào thực đơn và Xóa món ăn khỏi thực đơn."),
        ("Module Danh sách Mua sắm:", "Xem danh sách mua sắm, Thêm nguyên liệu cần mua, Đánh dấu đã mua nguyên liệu (hoàn thành) và Xóa nguyên liệu mua sắm."),
        ("Module Tư vấn sức khỏe (AI Coach Chatbot):", "Tư vấn dinh dưỡng qua Chatbot và Xem lịch sử trò chuyện."),
        ("Module Quản trị hệ thống (Admin Dashboard):", "Xem Dashboard thống kê hệ thống, Xem danh sách người dùng, Tìm kiếm người dùng, Khóa tài khoản người dùng, Mở khóa tài khoản người dùng, Phân quyền tài khoản (Nâng cấp/Hạ cấp USER/ADMIN), Xem toàn bộ lịch sử quét hệ thống, Tìm kiếm lịch sử quét, Xuất báo cáo thống kê quét thực phẩm, Xem danh sách thư viện thực phẩm, Tìm kiếm thực phẩm thư viện, Thêm thực phẩm vào thư viện, Chỉnh sửa thực phẩm thư viện, Xóa thực phẩm khỏi thư viện và Điều chỉnh cấu hình hệ thống.")
    ]
    
    for title, desc in features:
        p_feat = doc.add_paragraph(style='List Bullet')
        p_feat.paragraph_format.space_after = Pt(3)
        run_title = p_feat.add_run(title + " ")
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(6, 95, 70)
        p_feat.add_run(desc)

    # ----------------------------------------------------
    # SECTION 3: PHÂN QUYỀN NGƯỜI DÙNG & USE CASE CHI TIẾT
    # ----------------------------------------------------
    add_heading_styled(doc, "3. PHÂN QUYỀN VÀ MÔ TẢ CHI TIẾT SƠ ĐỒ USE CASE", 1)
    
    p = doc.add_paragraph()
    p.add_run("Hệ thống xác định rõ 3 đối tượng tác nhân (Actors) chính và các Use Case chi tiết cụ thể, không sử dụng các khái niệm quản trị chung chung:")
    
    # Table of Detailed Use Cases
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tác nhân (Actor)'
    hdr_cells[1].text = 'Vai trò hệ thống'
    hdr_cells[2].text = 'Các Use Case cụ thể (Chức năng chi tiết)'
    
    # Format header
    for cell in hdr_cells:
        set_cell_background(cell, "065F46")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                
    role_data = [
        (
            "Khách vãng lai\n(Guest)",
            "Người dùng chưa có tài khoản hoặc chưa đăng nhập vào hệ thống.",
            "- Đăng ký tài khoản mới\n- Đăng nhập hệ thống\n- Tải ảnh thực phẩm từ thiết bị\n- Chụp ảnh thực phẩm bằng camera\n- Nhận diện thực phẩm qua AI\n- Xem kết quả phân tích dinh dưỡng"
        ),
        (
            "Người dùng thành viên\n(Member User)",
            "Người dùng thông thường đã đăng nhập thành công vào ứng dụng.",
            "- Đăng xuất tài khoản\n- Xem thông tin cá nhân\n- Chỉnh sửa hồ sơ cá nhân\n- Thay đổi mật khẩu\n- Quét thực phẩm & Lưu lịch sử quét dinh dưỡng\n- Xem lịch sử quét cá nhân\n- Xóa lịch sử quét cá nhân\n- Ghi nhật ký cân nặng\n- Xem biểu đồ cân nặng & BMI\n- Ghi nhận lượng nước uống\n- Ghi nhận thời gian vận động\n- Xem biểu đồ dinh dưỡng tuần\n- Lập thực đơn bữa ăn\n- Xem thực đơn bữa ăn\n- Thêm món ăn vào thực đơn\n- Xóa món ăn khỏi thực đơn\n- Xem danh sách mua sắm\n- Thêm nguyên liệu cần mua\n- Đánh dấu đã mua nguyên liệu\n- Xóa nguyên liệu mua sắm\n- Tư vấn dinh dưỡng qua Chatbot\n- Xem lịch sử trò chuyện"
        ),
        (
            "Quản trị viên\n(Administrator)",
            "Người quản lý có đặc quyền cao nhất trong hệ thống.",
            "- Đăng nhập hệ thống\n- Đăng xuất tài khoản\n- Xem Dashboard thống kê hệ thống\n- Xem danh sách người dùng\n- Tìm kiếm người dùng\n- Khóa tài khoản người dùng\n- Mở khóa tài khoản người dùng\n- Phân quyền tài khoản (Nâng cấp/Hạ cấp USER/ADMIN)\n- Xem toàn bộ lịch sử quét hệ thống\n- Tìm kiếm lịch sử quét\n- Xuất báo cáo thống kê quét thực phẩm\n- Xem danh sách thư viện thực phẩm\n- Tìm kiếm thực phẩm thư viện\n- Thêm thực phẩm vào thư viện\n- Chỉnh sửa thực phẩm thư viện\n- Xóa thực phẩm khỏi thư viện\n- Điều chỉnh cấu hình hệ thống"
        )
    ]
    
    for r_name, r_scope, r_funcs in role_data:
        row = table.add_row()
        cells = row.cells
        cells[0].text = r_name
        cells[1].text = r_scope
        cells[2].text = r_funcs
        
        # Apply zebra coloring and margins
        for i, cell in enumerate(cells):
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(6, 95, 70)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    
    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # ----------------------------------------------------
    # SECTION 4: ACTIVITY DIAGRAM
    # ----------------------------------------------------
    add_heading_styled(doc, "4. QUY TRÌNH NHẬN DIỆN VÀ BIỂU ĐỒ HOẠT ĐỘNG (ACTIVITY DIAGRAM)", 1)
    
    p = doc.add_paragraph()
    p.add_run("Quy trình chính của hệ thống là ")
    p.add_run("Nhận diện hình ảnh bằng AI (Tải/Chụp ảnh -> AI xử lý -> Trả kết quả)").bold = True
    p.add_run(". Dưới đây mô tả chi tiết luồng xử lý của quy trình này qua biểu đồ hoạt động:")

    add_heading_styled(doc, "4.1 Các thành phần trong quy trình nhận diện", 2)
    
    components_flow = [
        ("Điểm bắt đầu (Bắt đầu): ", "Người dùng truy cập màn hình quét, lựa chọn tải ảnh thực phẩm từ thiết bị lên hoặc kích hoạt Camera chụp trực tiếp."),
        ("Bước xử lý phía Client: ", "Mã hóa hình ảnh sang định dạng base64/buffer, đính kèm mã xác thực JWT trong Header Authorization (nếu đã đăng nhập) và gửi yêu cầu POST HTTP tới cổng Node.js Gateway."),
        ("Bước tiền xử lý & Tối ưu hóa (Node.js Gateway): ", "Nhận luồng ảnh buffer. Sử dụng thư viện Jimp để kiểm tra kích thước. Nếu chiều rộng lớn hơn 800px, tiến hành co ảnh về 800px (giữ nguyên tỷ lệ) và nén chất lượng JPEG xuống 75% để giảm tải băng thông truyền tải lên mô hình AI."),
        ("Bước gọi phân tích AI (Gemini Service): ", "Gateway truyền ảnh cùng System Prompt định hình kết quả JSON sang Gemini API Vision. Gemini phân tích ảnh và kiểm tra điều kiện."),
        ("Điều kiện rẽ nhánh 1 (Có phải thực phẩm?): ", "AI đánh giá xem ảnh có chứa thực phẩm hay không. Nếu không, trả về isFood = false. Nếu có, trích xuất dữ liệu dinh dưỡng (Calo, Đạm, Tinh bột, Béo, Healthy Score) và trả về isFood = true."),
        ("Điều kiện rẽ nhánh 2 (Kiểm tra đăng nhập): ", "Gateway nhận dữ liệu từ Gemini. Nếu isFood = true và người dùng đã đăng nhập, Gateway gửi request ghi lịch sử sang Spring Boot backend để lưu trữ. Nếu chưa đăng nhập hoặc ảnh không phải thực phẩm, bỏ qua bước này."),
        ("Bước lưu cơ sở dữ liệu (Spring Boot): ", "Nhận thông tin dinh dưỡng từ Gateway, tạo bản ghi ScanHistory liên kết với ID người dùng và lưu vào MySQL database, trả về phản hồi thành công."),
        ("Kết quả đầu ra (Output) & Điểm kết thúc: ", "Gateway trả kết quả phân tích về Client. Client hiển thị kết quả trực quan trên màn hình (vẽ biểu đồ tròn chất đa lượng và điểm sức khỏe) hoặc hiển thị thông báo lỗi (nếu không phải thực phẩm). Quy trình kết thúc thành công hoặc kết thúc lỗi.")
    ]

    for title, desc in components_flow:
        p_comp = doc.add_paragraph(style='List Bullet')
        p_comp.paragraph_format.space_after = Pt(3)
        run_t = p_comp.add_run(title)
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(15, 118, 110)
        p_comp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # ----------------------------------------------------
    # SECTION 5: CƠ SỞ DỮ LIỆU VÀ CẤU TRÚC BẢNG (DATABASE SCHEMA)
    # ----------------------------------------------------
    add_heading_styled(doc, "5. CƠ SỞ DỮ LIỆU VÀ CẤU TRÚC BẢNG", 1)
    p = doc.add_paragraph()
    p.add_run("Cơ sở dữ liệu chi tiết của hệ thống AI Food Scanner trên MySQL được thiết kế chặt chẽ phục vụ cho việc lưu trữ các chức năng cụ thể:")

    # Table: Users
    add_heading_styled(doc, "5.1 Bảng: users (Thông tin tài khoản)", 2)
    tbl_users = doc.add_table(rows=1, cols=4)
    tbl_users.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_users)
    hdr_u = tbl_users.rows[0].cells
    hdr_u[0].text, hdr_u[1].text, hdr_u[2].text, hdr_u[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr_u:
        set_cell_background(cell, "0f766e")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    u_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính tài khoản"),
        ("username", "VARCHAR(255)", "UNIQUE, NOT NULL", "Tên đăng nhập"),
        ("password", "VARCHAR(255)", "NOT NULL", "Mật khẩu BCrypt"),
        ("email", "VARCHAR(255)", "UNIQUE, NOT NULL", "Email liên hệ"),
        ("full_name", "VARCHAR(255)", "NULL", "Họ tên đầy đủ"),
        ("role", "VARCHAR(50)", "NOT NULL", "Vai trò: ROLE_USER, ROLE_ADMIN"),
        ("enabled", "BOOLEAN", "NOT NULL, DEFAULT TRUE", "Trạng thái tài khoản (Khóa/Mở khóa)")
    ]
    for field, dtype, const, desc in u_fields:
        r = tbl_users.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_before = Pt(10)

    # Table: scan_histories
    add_heading_styled(doc, "5.2 Bảng: scan_histories (Lịch sử quét thực phẩm)", 2)
    tbl_scans = doc.add_table(rows=1, cols=4)
    tbl_scans.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_scans)
    hdr_s = tbl_scans.rows[0].cells
    hdr_s[0].text, hdr_s[1].text, hdr_s[2].text, hdr_s[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr_s:
        set_cell_background(cell, "0f766e")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    s_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính lượt quét"),
        ("user_id", "BIGINT", "FOREIGN KEY -> users(id)", "Liên kết với người dùng thực hiện quét"),
        ("food_name", "VARCHAR(255)", "NOT NULL", "Tên món ăn do AI nhận dạng"),
        ("calories", "DOUBLE", "NOT NULL", "Lượng calo"),
        ("protein", "DOUBLE", "NOT NULL", "Lượng protein (gram)"),
        ("carbs", "DOUBLE", "NOT NULL", "Lượng carbohydrate (gram)"),
        ("fat", "DOUBLE", "NOT NULL", "Lượng chất béo (gram)"),
        ("healthy_score", "INTEGER", "NOT NULL", "Điểm số sức khỏe thực phẩm (0 - 100)"),
        ("raw_json_result", "TEXT", "NULL", "Dữ liệu JSON gốc từ Gemini API"),
        ("created_at", "DATETIME", "NOT NULL", "Ngày giờ quét thực phẩm")
    ]
    for field, dtype, const, desc in s_fields:
        r = tbl_scans.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_before = Pt(10)

    # Save document
    filename = "Bao_cao_He_thong_AI_Food_Scanner.docx"
    doc.save(filename)
    print(f"Detailed Report updated successfully: {filename}")

if __name__ == '__main__':
    main()
