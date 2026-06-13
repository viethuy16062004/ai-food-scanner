import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="A0AEC0"/>'
        '<w:left w:val="none"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="718096"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>'
    )
    tblBorders = parse_xml(borders_xml)
    tblPr.append(tblBorders)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        run.font.size = Pt(13.5)
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_custom_paragraph(doc, text, is_italic=False, space_after=6, list_bullet=False):
    style = 'List Bullet' if list_bullet else 'Normal'
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    
    if not list_bullet:
        p.paragraph_format.first_line_indent = Inches(0.5)
        
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(30, 41, 59)
    run.italic = is_italic
    return p

def main():
    doc = docx.Document()
    
    # Page Margins (Top/Bottom = 2.0 cm, Left = 3.0 cm, Right = 2.0 cm)
    # 1 cm = 0.3937 inches
    for section in doc.sections:
        section.top_margin = Inches(0.787)
        section.bottom_margin = Inches(0.787)
        section.left_margin = Inches(1.181)
        section.right_margin = Inches(0.787)
        
    # Styles configuration
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(30, 41, 59)

    # =========================================================================
    # CHƯƠNG 2: NỘI DUNG THỰC TẬP
    # =========================================================================
    p_chap = doc.add_paragraph()
    p_chap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_chap.paragraph_format.space_before = Pt(24)
    p_chap.paragraph_format.space_after = Pt(18)
    run_chap = p_chap.add_run("CHƯƠNG 2: NỘI DUNG THỰC TẬP")
    run_chap.font.name = 'Times New Roman'
    run_chap.font.size = Pt(16)
    run_chap.bold = True
    run_chap.font.color.rgb = RGBColor(15, 23, 42)

    # -------------------------------------------------------------------------
    # 2.1. Tìm hiểu cơ sở lý thuyết và công nghệ áp dụng
    # -------------------------------------------------------------------------
    add_heading_styled(doc, "2.1. Tìm hiểu cơ sở lý thuyết và công nghệ áp dụng", 1)
    
    add_custom_paragraph(doc, 
        "Trong quá trình thực tập, việc nghiên cứu các cơ sở lý thuyết vững chắc và lựa chọn công nghệ "
        "phù hợp đóng vai trò quyết định đến sự thành công của hệ thống. Nhóm đã thực hiện khảo sát các "
        "mô hình quản lý dự án hiện đại, tìm hiểu phương pháp phát triển phần mềm linh hoạt, nghiên cứu "
        "cơ chế phân quyền chặt chẽ cũng như lựa chọn các công nghệ lập trình tiên tiến nhất để tối ưu hóa "
        "hiệu năng và trải nghiệm người dùng đối với ứng dụng AI NutriScan.")

    # 2.1.1. Công cụ quản lý công việc
    add_heading_styled(doc, "2.1.1. Công cụ quản lý công việc", 2)
    
    add_custom_paragraph(doc, 
        "Để đảm bảo sự phối hợp nhịp nhàng giữa các thành viên và theo dõi tiến độ phát triển một cách chính xác, "
        "nhóm đã áp dụng các công cụ quản lý công việc tiêu chuẩn công nghiệp:")
        
    add_custom_paragraph(doc, 
        "Git & GitHub: Được sử dụng để quản lý mã nguồn phân tán. Nhóm áp dụng mô hình Git Flow, "
        "mỗi chức năng mới sẽ được phát triển trên một nhánh riêng (feature branch) và chỉ được tích hợp "
        "vào nhánh chính (main) sau khi trải qua quá trình Pull Request và Review mã nguồn nghiêm ngặt.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Trello/Jira Board: Công cụ trực quan hóa tiến độ công việc dựa trên bảng Kanban. Các nhiệm vụ "
        "được chia nhỏ và phân loại vào các cột: To Do (Cần làm), In Progress (Đang thực hiện), "
        "Review (Đang đánh giá chất lượng code) và Done (Đã hoàn thành), giúp cả nhóm nắm bắt tức thời trạng thái dự án.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Slack & Discord: Sử dụng làm kênh trao đổi thông tin chính thức. Tích hợp webhook tự động báo cáo "
        "mỗi khi có commit mới hoặc thay đổi trạng thái task trên bảng quản lý công việc.", 
        list_bullet=True)

    # 2.1.2. Phương pháp Scrum
    add_heading_styled(doc, "2.1.2. Phương pháp Scrum", 2)
    
    add_custom_paragraph(doc, 
        "Dự án được vận hành theo mô hình phát triển phần mềm Agile thông qua phương pháp Scrum. "
        "Quy trình triển khai bao gồm các vai trò, tài liệu và cuộc họp đặc thù giúp tăng tính linh hoạt "
        "và khả năng thích ứng cao trước các thay đổi kỹ thuật:")
        
    add_custom_paragraph(doc, 
        "Vai trò trong nhóm: Trần Việt Huy đảm nhận vai trò Trưởng nhóm (Scrum Master & Lead Developer) điều phối "
        "các hoạt động kỹ thuật, trong khi các thành viên còn lại đóng vai trò là Nhóm phát triển (Development Team) "
        "cùng tham gia xây dựng và kiểm thử hệ thống.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Chu kỳ Sprint: Mỗi Sprint kéo dài 2 tuần. Trước khi bắt đầu mỗi Sprint, nhóm tổ chức họp Sprint Planning "
        "để lựa chọn các yêu cầu từ Product Backlog đưa vào Sprint Backlog dựa trên năng suất và độ ưu tiên.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Họp đứng hàng ngày (Daily Standup): Diễn ra trong khoảng 10-15 phút vào đầu ngày làm việc. Mỗi thành viên "
        "trả lời ba câu hỏi cốt lõi: Hôm qua đã làm gì? Hôm nay sẽ làm gì? Có gặp khó khăn hay trở ngại gì không?", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Đánh giá và cải tiến (Sprint Review & Retrospective): Kết thúc mỗi Sprint, nhóm demo sản phẩm có thể chạy "
        "được (Increment) cho giảng viên/người hướng dẫn để lấy phản hồi, đồng thời họp rút kinh nghiệm nhằm cải tiến "
        "quy trình làm việc cho Sprint tiếp theo.", 
        list_bullet=True)

    # 2.1.3. Mô hình phân quyền (RBAC)
    add_heading_styled(doc, "2.1.3. Mô hình phân quyền (RBAC)", 2)
    
    add_custom_paragraph(doc, 
        "Mô hình phân quyền dựa trên vai trò (Role-Based Access Control - RBAC) là cơ chế bảo mật "
        "cốt lõi được tích hợp trong hệ thống để quản lý quyền hạn truy cập của người dùng. Hệ thống chia "
        "làm 3 nhóm đối tượng cụ thể với các quyền hạn được phân cấp rõ ràng:")
        
    add_custom_paragraph(doc, 
        "Khách vãng lai (Guest): Là người dùng chưa đăng nhập hệ thống. Chỉ được phép xem trang giới thiệu chung, "
        "đăng ký tài khoản, đăng nhập và sử dụng thử chức năng quét thực phẩm AI ở mức cơ bản (không lưu lịch sử).", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Người dùng thành viên (Member User): Là người dùng thông thường đã đăng nhập. Có quyền sử dụng đầy đủ các "
        "chức năng cá nhân hóa bao gồm: lưu lịch sử quét, ghi nhật ký chỉ số cơ thể, lập kế hoạch bữa ăn, theo dõi "
        "danh sách mua sắm, chỉnh sửa thông tin hồ sơ và trò chuyện tư vấn dinh dưỡng với AI Coach Chatbot.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Quản trị viên (Administrator): Có quyền quản lý cao nhất. Admin có quyền truy cập vào dashboard thống kê "
        "hệ thống, quản lý danh sách người dùng (khóa/mở khóa tài khoản, nâng/hạ quyền hạn), kiểm tra lịch sử quét "
        "toàn hệ thống, cập nhật danh mục thư viện thực phẩm dùng chung và điều chỉnh các tham số cấu hình hệ thống.", 
        list_bullet=True)

    # Table for RBAC Summary
    table_rbac = doc.add_table(rows=1, cols=3)
    table_rbac.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_rbac)
    
    hdr_rbac = table_rbac.rows[0].cells
    hdr_rbac[0].text = 'Vai trò (Role)'
    hdr_rbac[1].text = 'Phạm vi truy cập'
    hdr_rbac[2].text = 'Các chức năng chính được phép thực hiện'
    
    for cell in hdr_rbac:
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
                
    rbac_data = [
        ("Guest", "Công khai (Public)", "Đăng ký, Đăng nhập, Trải nghiệm quét thực phẩm AI (không lưu dữ liệu)."),
        ("Member User", "Cá nhân (Private - Auth)", "Quét thực phẩm và lưu lịch sử, Nhật ký sức khỏe, Lập kế hoạch bữa ăn, Danh sách mua sắm, Chatbot tư vấn."),
        ("Administrator", "Quản trị (Admin - Strict Auth)", "Dashboard thống kê, Quản lý tài khoản (Khóa/Mở khóa/Phân quyền), Quản lý Thư viện thực phẩm, Cấu hình hệ thống.")
    ]
    
    for role, scope, funcs in rbac_data:
        row = table_rbac.add_row()
        row.cells[0].text = role
        row.cells[1].text = scope
        row.cells[2].text = funcs
        for i, cell in enumerate(row.cells):
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # 2.1.4. Công nghệ triển khai
    add_heading_styled(doc, "2.1.4. Công nghệ triển khai", 2)
    
    add_custom_paragraph(doc, 
        "Hệ thống ứng dụng kiến trúc phân tầng hiện đại nhằm tối ưu hóa hiệu năng, tính bảo mật "
        "và khả năng mở rộng. Các công nghệ cốt lõi bao gồm:")
        
    add_custom_paragraph(doc, 
        "Frontend: Xây dựng trên nền tảng Next.js 15 (React 19) kết hợp ngôn ngữ TypeScript. "
        "Sử dụng thư viện Tailwind CSS để thiết kế giao diện responsive theo phong cách Modern Glassmorphism. "
        "Trạng thái ứng dụng được quản lý hiệu quả qua thư viện Zustand và hiệu ứng chuyển động mượt mà "
        "được thực hiện bởi Framer Motion.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Backend Core Service: Sử dụng framework Spring Boot 3.3.0 (Java 21). Đảm nhận việc xử lý "
        "toàn bộ logic nghiệp vụ, quản lý dữ liệu người dùng, lưu lịch sử, kế hoạch ăn uống và bảo mật "
        "thông qua Spring Security kết hợp với mã hóa JSON Web Token (JWT) Stateless.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Node.js Gemini AI Gateway: Một cổng API trung gian được xây dựng trên Express.js, đảm nhận nhiệm vụ "
        "tiền xử lý hình ảnh (sử dụng thư viện Jimp để co kích thước về tối đa 800px và nén chất lượng JPEG xuống 75%) "
        "và thực hiện kết nối bảo mật tới Google Gemini API để phân tích hình ảnh thực phẩm mà không lộ API key về phía Client.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Cơ sở dữ liệu: Hệ quản trị cơ sở dữ liệu quan hệ PostgreSQL được chọn làm nơi lưu trữ dữ liệu "
        "nhờ tính tin cậy cao, hỗ trợ tốt các ràng buộc khóa ngoại phức tạp và khả năng lập chỉ mục (index) tối ưu.", 
        list_bullet=True)

    # -------------------------------------------------------------------------
    # 2.2. Phân tích và thiết kế hệ thống
    # -------------------------------------------------------------------------
    add_heading_styled(doc, "2.2. Phân tích và thiết kế hệ thống", 1)
    
    add_custom_paragraph(doc, 
        "Để hiện thực hóa các yêu cầu nghiệp vụ thành một sản phẩm phần mềm ổn định và có cấu trúc "
        "chặt chẽ, nhóm đã tiến hành phân tích chi tiết bài toán thực tế, làm rõ các yêu cầu hệ thống, "
        "thiết kế các sơ đồ mô hình hóa và xây dựng lược đồ cơ sở dữ liệu chi tiết.")

    # 2.2.1. Phân tích bài toán
    add_heading_styled(doc, "2.2.1. Phân tích bài toán", 2)
    
    add_custom_paragraph(doc, 
        "Trong bối cảnh hiện đại, tỷ lệ mắc các bệnh liên quan đến chế độ ăn uống không khoa học như béo phì, "
        "tiểu đường, tim mạch đang ngày càng gia tăng. Việc tự theo dõi hàm lượng dinh dưỡng tiêu thụ hàng ngày "
        "là một nhu cầu thiết thực, tuy nhiên quy trình này thường gặp phải các rào cản lớn:")
        
    add_custom_paragraph(doc, 
        "Khó khăn trong việc xác định định lượng và calo của món ăn: Người dùng thông thường không thể biết "
        "chính xác một đĩa cơm tấm hay bát phở chứa bao nhiêu calo, protein, carbs hay fat nếu không có công cụ chuyên dụng.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Quy trình nhập liệu thủ công tốn thời gian: Việc tra cứu thủ công và nhập thông số từng nguyên liệu "
        "vào sổ tay hoặc các ứng dụng cũ làm người dùng nhanh nản lòng và từ bỏ thói quen ghi chép dinh dưỡng.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Giải pháp của hệ thống AI NutriScan là ứng dụng trí tuệ nhân tạo thị giác máy tính thông qua mô hình "
        "Gemini Vision API. Người dùng chỉ cần chụp ảnh món ăn hoặc nhãn dinh dưỡng bao bì sản phẩm, hệ thống sẽ "
        "tự động nhận dạng, phân tích định lượng calo và các chất đa lượng (protein, fat, carbs) chỉ trong 1-2 giây. "
        "Kết hợp với các module theo dõi sức khỏe, hệ thống mang đến một giải pháp chăm sóc sức khỏe toàn diện và liền mạch.", 
        list_bullet=True)

    # 2.2.2. Yêu cầu hệ thống
    add_heading_styled(doc, "2.2.2. Yêu cầu hệ thống", 2)
    
    add_custom_paragraph(doc, 
        "Yêu cầu hệ thống được chia làm hai nhóm chính bao gồm yêu cầu chức năng (những gì hệ thống phải làm) "
        "và yêu cầu phi chức năng (các tiêu chuẩn chất lượng vận hành):")
        
    add_custom_paragraph(doc, 
        "Yêu cầu chức năng (Functional Requirements):", is_italic=True)
        
    add_custom_paragraph(doc, 
        "Module Tài khoản: Đăng ký, đăng nhập (JWT), cập nhật thông tin cá nhân (chiều cao, mục tiêu calo), đổi mật khẩu.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Module Quét AI: Cho phép tải ảnh thực phẩm hoặc chụp camera trực tiếp. Trích xuất thông số dinh dưỡng từ món ăn "
        "hoặc nhận diện bảng Nutrition Facts trên bao bì sản phẩm qua OCR. Lưu lịch sử quét đối với tài khoản đã đăng nhập.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Module Nhật ký sức khỏe: Ghi nhận chỉ số cân nặng, lượng nước uống, thời gian vận động hàng ngày và vẽ biểu đồ theo dõi.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Module Kế hoạch bữa ăn & Danh sách mua sắm: Lập thực đơn dinh dưỡng theo ngày/tuần, thêm món ăn vào các bữa ăn "
        "(Sáng, Trưa, Tối, Phụ). Tự động tổng hợp danh sách nguyên liệu cần mua sắm và cho phép tick hoàn thành.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Module Quản trị (Admin): Dashboard thống kê tăng trưởng, quản lý tài khoản người dùng, xem toàn bộ lịch sử quét, "
        "thêm/sửa/xóa thư viện thực phẩm mẫu.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Yêu cầu phi chức năng (Non-functional Requirements):", is_italic=True)
        
    add_custom_paragraph(doc, 
        "Hiệu năng (Performance): Thời gian phản hồi phân tích hình ảnh AI dưới 3 giây đối với các ảnh đã được tối ưu hóa dung lượng.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Bảo mật (Security): Mã hóa mật khẩu người dùng bằng thuật toán BCrypt. Toàn bộ các API nhạy cảm phải được bảo vệ "
        "bằng bộ lọc Spring Security JWT và chặn tấn công Cross-Origin Resource Sharing (CORS).", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Độ tương thích (Usability): Giao diện hiển thị trực quan, mượt mà trên cả trình duyệt máy tính và thiết bị di động (Responsive).", 
        list_bullet=True)

    # 2.2.3. Mô hình hóa hệ thống
    add_heading_styled(doc, "2.2.3. Mô hình hóa hệ thống", 2)
    
    add_custom_paragraph(doc, 
        "Để làm rõ kiến trúc tương tác và các luồng xử lý chính trong ứng dụng, nhóm đã tiến hành xây dựng sơ đồ Use Case "
        "tổng quát và sơ đồ Hoạt động (Activity Diagram) mô tả chi tiết quy trình nhận diện thực phẩm bằng công nghệ AI:")

    add_custom_paragraph(doc, 
        "Sơ đồ Use Case tổng quát của hệ thống mô tả mối quan hệ giữa 3 tác nhân (Guest, Member, Admin) "
        "và các chức năng tương ứng:")
        
    # Embed Use Case Diagram
    use_case_path = r"c:\CNPM\ai-food-scanner\use_case_diagram.png"
    if os.path.exists(use_case_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(use_case_path, width=Inches(5.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run("Hình 2.1. Sơ đồ Use Case chi tiết của hệ thống AI NutriScan")
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10.5)
        run_cap.italic = True
    else:
        add_custom_paragraph(doc, "[Hình 2.1: Sơ đồ Use Case hệ thống - file use_case_diagram.png không tìm thấy]", is_italic=True)

    add_custom_paragraph(doc, 
        "Quy trình xử lý cốt lõi của hệ thống là chức năng nhận diện hình ảnh món ăn và lưu trữ lịch sử dinh dưỡng. "
        "Luồng hoạt động này được mô tả chi tiết qua Sơ đồ Hoạt động (Activity Diagram) dưới đây:")

    # Embed Activity Diagram
    activity_path = r"c:\CNPM\ai-food-scanner\activity_diagram.png"
    if os.path.exists(activity_path):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.paragraph_format.space_before = Pt(8)
        p_img2.paragraph_format.space_after = Pt(4)
        run_img2 = p_img2.add_run()
        run_img2.add_picture(activity_path, width=Inches(5.5))
        
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap2.paragraph_format.space_after = Pt(12)
        run_cap2 = p_cap2.add_run("Hình 2.2. Sơ đồ Hoạt động quy trình quét và phân tích dinh dưỡng bằng AI")
        run_cap2.font.name = 'Times New Roman'
        run_cap2.font.size = Pt(10.5)
        run_cap2.italic = True
    else:
        add_custom_paragraph(doc, "[Hình 2.2: Sơ đồ Hoạt động - file activity_diagram.png không tìm thấy]", is_italic=True)

    add_custom_paragraph(doc, 
        "Theo sơ đồ hoạt động, khi người dùng tải hoặc chụp ảnh, Client gửi ảnh thô lên Node.js Gateway. "
        "Tại đây, thư viện Jimp sẽ tự động thực hiện nén và co kích thước ảnh trước khi chuyển tiếp sang Gemini API. "
        "Mô hình AI nhận diện và kiểm tra xem ảnh có chứa thực phẩm hay không. Nếu có, dữ liệu JSON chứa các thông số dinh dưỡng "
        "được trích xuất và trả về. Gateway sau đó sẽ kiểm tra trạng thái đăng nhập của người dùng: nếu đã đăng nhập, "
        "một yêu cầu POST sẽ được gửi sang Spring Boot API để lưu dữ liệu vào MySQL Database, đảm bảo tính nhất quán dữ liệu.")

    # 2.2.4. Thiết kế cơ sở dữ liệu
    add_heading_styled(doc, "2.2.4. Thiết kế cơ sở dữ liệu", 2)
    
    add_custom_paragraph(doc, 
        "Cơ sở dữ liệu của hệ thống được thiết kế theo mô hình quan hệ chuẩn hóa để tránh dư thừa dữ liệu "
        "và đảm bảo tính toàn vẹn tham chiếu. Dưới đây là cấu trúc chi tiết của các bảng cơ sở dữ liệu chính:")

    # Table 1: users
    add_heading_styled(doc, "Bảng: users (Thông tin tài khoản)", 3)
    t_users = doc.add_table(rows=1, cols=4)
    t_users.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_users)
    hdr = t_users.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr:
        set_cell_background(cell, "334155")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    u_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính tài khoản"),
        ("username", "VARCHAR(255)", "UNIQUE, NOT NULL", "Tên đăng nhập"),
        ("password", "VARCHAR(255)", "NOT NULL", "Mật khẩu mã hóa BCrypt"),
        ("email", "VARCHAR(255)", "UNIQUE, NOT NULL", "Địa chỉ email người dùng"),
        ("full_name", "VARCHAR(255)", "NULL", "Họ tên đầy đủ"),
        ("role", "VARCHAR(50)", "NOT NULL", "Vai trò: ROLE_USER, ROLE_ADMIN"),
        ("enabled", "BOOLEAN", "NOT NULL, DEFAULT TRUE", "Trạng thái tài khoản (Khóa/Mở khóa)")
    ]
    for field, dtype, const, desc in u_fields:
        r = t_users.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True
                
    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # Table 2: scan_histories
    add_heading_styled(doc, "Bảng: scan_histories (Lịch sử quét thực phẩm)", 3)
    t_scans = doc.add_table(rows=1, cols=4)
    t_scans.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_scans)
    hdr_s = t_scans.rows[0].cells
    hdr_s[0].text, hdr_s[1].text, hdr_s[2].text, hdr_s[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr_s:
        set_cell_background(cell, "334155")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    s_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính lượt quét"),
        ("user_id", "BIGINT", "FOREIGN KEY -> users(id)", "Liên kết với người dùng thực hiện quét"),
        ("food_name", "VARCHAR(255)", "NOT NULL", "Tên món ăn do AI nhận dạng"),
        ("calories", "DOUBLE", "NOT NULL", "Lượng calo phân tích"),
        ("protein", "DOUBLE", "NOT NULL", "Hàm lượng chất đạm (gram)"),
        ("carbs", "DOUBLE", "NOT NULL", "Hàm lượng tinh bột (gram)"),
        ("fat", "DOUBLE", "NOT NULL", "Hàm lượng chất béo (gram)"),
        ("healthy_score", "INTEGER", "NOT NULL", "Điểm số sức khỏe thực phẩm (0 - 100)"),
        ("raw_json_result", "TEXT", "NULL", "Dữ liệu JSON gốc từ Gemini API"),
        ("created_at", "DATETIME", "NOT NULL", "Ngày giờ quét thực phẩm")
    ]
    for field, dtype, const, desc in s_fields:
        r = t_scans.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # Table 3: health_logs
    add_heading_styled(doc, "Bảng: health_logs (Nhật ký chỉ số cơ thể hàng ngày)", 3)
    t_health = doc.add_table(rows=1, cols=4)
    t_health.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_health)
    hdr_h = t_health.rows[0].cells
    hdr_h[0].text, hdr_h[1].text, hdr_h[2].text, hdr_h[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr_h:
        set_cell_background(cell, "334155")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    h_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính nhật ký sức khỏe"),
        ("user_id", "BIGINT", "FOREIGN KEY -> users(id)", "Liên kết với người dùng"),
        ("weight", "DOUBLE", "NULL", "Chỉ số cân nặng (kg)"),
        ("bmi", "DOUBLE", "NULL", "Chỉ số khối cơ thể"),
        ("body_fat_percent", "DOUBLE", "NULL", "Tỷ lệ phần trăm mỡ cơ thể"),
        ("water_intake_ml", "DOUBLE", "NULL", "Lượng nước uống ghi nhận (ml)"),
        ("active_minutes", "DOUBLE", "NULL", "Thời gian vận động (phút)"),
        ("log_date", "DATE", "NOT NULL", "Ngày ghi nhận chỉ số"),
        ("created_at", "DATETIME", "NULL", "Thời điểm khởi tạo bản ghi")
    ]
    for field, dtype, const, desc in h_fields:
        r = t_health.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # Table 4: meal_plans
    add_heading_styled(doc, "Bảng: meal_plans (Kế hoạch bữa ăn ngày)", 3)
    t_plans = doc.add_table(rows=1, cols=4)
    t_plans.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_plans)
    hdr_p = t_plans.rows[0].cells
    hdr_p[0].text, hdr_p[1].text, hdr_p[2].text, hdr_p[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr_p:
        set_cell_background(cell, "334155")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    p_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính kế hoạch ngày"),
        ("user_id", "BIGINT", "FOREIGN KEY -> users(id)", "Liên kết tài khoản người dùng"),
        ("goal_calories", "DOUBLE", "NULL", "Mục tiêu calo trong ngày"),
        ("goal_protein", "DOUBLE", "NULL", "Mục tiêu chất đạm nạp vào (g)"),
        ("goal_carbs", "DOUBLE", "NULL", "Mục tiêu chất tinh bột nạp vào (g)"),
        ("goal_fat", "DOUBLE", "NULL", "Mục tiêu chất béo nạp vào (g)"),
        ("dietary_goal", "VARCHAR(255)", "NULL", "Chế độ ăn kiêng mục tiêu"),
        ("plan_date", "DATE", "NOT NULL", "Ngày lập kế hoạch thực đơn"),
        ("created_at", "DATETIME", "NULL", "Thời điểm khởi tạo kế hoạch")
    ]
    for field, dtype, const, desc in p_fields:
        r = t_plans.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # Table 5: meal_plan_items
    add_heading_styled(doc, "Bảng: meal_plan_items (Món ăn chi tiết trong thực đơn)", 3)
    t_items = doc.add_table(rows=1, cols=4)
    t_items.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_items)
    hdr_it = t_items.rows[0].cells
    hdr_it[0].text, hdr_it[1].text, hdr_it[2].text, hdr_it[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr_it:
        set_cell_background(cell, "334155")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    it_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính món ăn thực đơn"),
        ("meal_plan_id", "BIGINT", "FOREIGN KEY -> meal_plans(id)", "Liên kết với kế hoạch ngày"),
        ("meal_type", "VARCHAR(50)", "NOT NULL", "Loại bữa ăn: SÁNG, TRƯA, CHIỀU, TỐI"),
        ("food_name", "VARCHAR(255)", "NOT NULL", "Tên món ăn"),
        ("calories", "DOUBLE", "NULL", "Hàm lượng calo"),
        ("protein", "DOUBLE", "NULL", "Hàm lượng chất đạm"),
        ("image_url", "VARCHAR(1000)", "NULL", "Đường dẫn hình ảnh món ăn")
    ]
    for field, dtype, const, desc in it_fields:
        r = t_items.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # Table 6: shopping_items
    add_heading_styled(doc, "Bảng: shopping_items (Danh sách mua sắm)", 3)
    t_shop = doc.add_table(rows=1, cols=4)
    t_shop.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_shop)
    hdr_sh = t_shop.rows[0].cells
    hdr_sh[0].text, hdr_sh[1].text, hdr_sh[2].text, hdr_sh[3].text = 'Trường (Field)', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'
    for cell in hdr_sh:
        set_cell_background(cell, "334155")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    sh_fields = [
        ("id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Khóa chính danh mục mua sắm"),
        ("user_id", "BIGINT", "FOREIGN KEY -> users(id)", "Liên kết với người dùng"),
        ("item_name", "VARCHAR(255)", "NOT NULL", "Tên nguyên liệu cần mua sắm"),
        ("store_category", "VARCHAR(255)", "NULL", "Địa điểm mua sắm (Siêu thị, chợ...)"),
        ("checked", "BOOLEAN", "NOT NULL, DEFAULT FALSE", "Trạng thái đã mua xong hay chưa"),
        ("created_at", "DATETIME", "NULL", "Thời điểm thêm vào danh sách")
    ]
    for field, dtype, const, desc in sh_fields:
        r = t_shop.add_row()
        r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text = field, dtype, const, desc
        for i, c in enumerate(r.cells):
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0:
                c.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # -------------------------------------------------------------------------
    # 2.3. Quá trình xây dựng và triển khai
    # -------------------------------------------------------------------------
    add_heading_styled(doc, "2.3. Quá trình xây dựng và triển khai", 1)
    
    add_custom_paragraph(doc, 
        "Sau khi hoàn thiện thiết kế hệ thống và lược đồ cơ sở dữ liệu, nhóm đã tiến hành thiết lập "
        "môi trường phát triển, tổ chức cấu trúc mã nguồn dự án khoa học và giải quyết các bài toán "
        "logic nghiệp vụ và bảo mật thông tin.")

    # 2.3.1. Môi trường phát triển
    add_heading_styled(doc, "2.3.1. Môi trường phát triển", 2)
    
    add_custom_paragraph(doc, 
        "Hệ thống được phát triển trên môi trường đồng bộ để đảm bảo mã nguồn hoạt động chính xác "
        "giữa tất cả các thành viên phát triển và khi triển khai môi trường production:")
        
    add_custom_paragraph(doc, 
        "Hệ điều hành: Windows 10/11 & macOS.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Công cụ lập trình (IDE): IntelliJ IDEA cho phần backend Java Spring Boot, Visual Studio Code "
        "cho phần frontend Next.js và backend Node.js Gateway.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Bộ thư viện phát triển (SDK/Runtime): Java Development Kit (JDK 21 LTS), Node.js v20.x, PostgreSQL 15.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Môi trường container hóa và triển khai: Docker & Docker Compose được sử dụng để đóng gói toàn bộ "
        "ứng dụng thành các container độc lập, chạy cục bộ dễ dàng. Pipeline CI/CD được cấu hình tự động bằng "
        "GitHub Actions để tự động hóa quá trình chạy kiểm thử đơn vị (Unit Tests), build Docker image và deploy "
        "trực tiếp lên Cloud Railway khi có thay đổi trên nhánh main.", 
        list_bullet=True)

    # 2.3.2. Cấu trúc mã nguồn
    add_heading_styled(doc, "2.3.2. Cấu trúc mã nguồn", 2)
    
    add_custom_paragraph(doc, 
        "Hệ thống được tổ chức theo cấu trúc thư mục Monorepo rõ ràng để dễ dàng quản lý đồng bộ cả 3 phần "
        "của hệ thống:")
        
    add_custom_paragraph(doc, 
        "/frontend: Chứa mã nguồn ứng dụng web xây dựng bằng Next.js. Mã nguồn được tổ chức theo mô hình App Router "
        "của Next.js 15, phân tách rõ ràng giữa các trang (/app), các component giao diện dùng chung (/components), "
        "quản lý trạng thái (/store) và định nghĩa kiểu dữ liệu (/types).", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "/backend (Spring Boot Core): Chứa mã nguồn xử lý logic nghiệp vụ và dữ liệu. Dự án Java tuân thủ "
        "kiến trúc phân tầng chuẩn (MVC/Domain-driven): controller (tiếp nhận request), service (xử lý logic nghiệp vụ), "
        "repository (tương tác CSDL JPA), entity (định nghĩa các thực thể CSDL), và security (cấu hình phân quyền Spring Security).", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "/backend/index.js (Node.js Gateway): Tệp tin định nghĩa dịch vụ cổng trung gian Express.js chịu trách nhiệm "
        "tiền xử lý ảnh (sử dụng thư viện Jimp) và kết nối với Google Gemini API Vision.", 
        list_bullet=True)

    # 2.3.3. Xử lý logic và bảo mật
    add_heading_styled(doc, "2.3.3. Xử lý logic và bảo mật", 2)
    
    add_custom_paragraph(doc, 
        "Vấn đề bảo mật thông tin và tính chính xác của thuật toán xử lý là ưu tiên hàng đầu trong quá trình xây dựng hệ thống:")
        
    add_custom_paragraph(doc, 
        "Bảo mật tài khoản và phân quyền: Mật khẩu người dùng được băm bảo mật bằng thuật toán BCrypt trước khi lưu "
        "vào cơ sở dữ liệu. Cơ chế xác thực sử dụng JSON Web Token (JWT) Stateless, mã hóa thông tin người dùng "
        "và vai trò thành chuỗi ký tự ký số gửi kèm trong Header Authorization cho mỗi request HTTP gửi lên server. "
        "Spring Security cấu hình bộ lọc (SecurityFilterChain) chặn mọi truy cập trái phép vào các API của Member "
        "và Admin.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Tiền xử lý hình ảnh tối ưu băng thông: Để hạn chế việc tải tệp tin dung lượng lớn lên đám mây làm chậm "
        "hệ thống, Node.js Gateway sử dụng Jimp để đọc luồng ảnh buffer từ client, tự động kiểm tra kích thước "
        "và thực hiện co ảnh về chiều rộng tối đa 800px (giữ nguyên tỷ lệ), đồng thời nén chất lượng ảnh JPEG xuống 75%. "
        "Thao tác này giúp giảm kích thước ảnh từ 4-5MB xuống dưới 200KB mà vẫn đảm bảo độ sắc nét cho AI phân tích.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Kỹ nghệ Prompt của Gemini API (Prompt Engineering): Để Gemini API trả về kết quả cấu trúc dữ liệu JSON "
        "chính xác, nhóm thiết lập một System Prompt chi tiết. Prompt này yêu cầu mô hình AI bắt buộc phải trả về định dạng "
        "JSON chuẩn với các trường cụ thể: isFood, foodName, calories, protein, carbs, fat, healthyScore và advice. "
        "Nếu hình ảnh không phải thực phẩm, mô hình AI sẽ trả về isFood: false giúp client hiển thị cảnh báo phù hợp.", 
        list_bullet=True)

    # -------------------------------------------------------------------------
    # 2.4. Kết quả công việc đạt được
    # -------------------------------------------------------------------------
    add_heading_styled(doc, "2.4. Kết quả công việc đạt được", 1)
    
    add_custom_paragraph(doc, 
        "Trải qua quá trình nghiên cứu, thiết kế và phát triển nghiêm túc, nhóm đã hoàn thành toàn diện "
        "hệ thống AI NutriScan với chất lượng hoàn thiện cao. Kết quả công việc đạt được thể hiện qua các "
        "phân hệ giao diện chức năng hoạt động ổn định:")

    # 2.4.1. Giao diện người dùng
    add_heading_styled(doc, "2.4.1. Giao diện người dùng", 2)
    
    add_custom_paragraph(doc, 
        "Phân hệ giao diện dành cho người dùng thành viên được thiết kế theo phong cách tối giản, hiện đại, "
        "chuyển động mượt mà và tương thích tốt trên các thiết bị di động:")
        
    add_custom_paragraph(doc, 
        "Màn hình Quét AI (Food AI Scanner): Hỗ trợ kéo thả ảnh hoặc mở camera chụp trực tiếp. Hiển thị kết quả "
        "quét trực quan bằng biểu đồ tròn phân tích chất đa lượng (protein, fat, carbs) và thanh tiến trình điểm số "
        "sức khỏe (Healthy Score). Hệ thống cũng tích hợp chế độ quét bao bì sản phẩm (Packaging OCR Scanner) đọc "
        "bảng thành phần dinh dưỡng và cảnh báo dị ứng.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Màn hình Dashboard & Nhật ký sức khỏe: Biểu diễn trực quan lượng Calo đã nạp so với mục tiêu ngày. "
        "Hiển thị các widget cập nhật cân nặng, lượng nước uống hàng ngày (ml) và số phút tập luyện thể thao.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Màn hình Kế hoạch ăn uống & Mua sắm: Người dùng dễ dàng lên thực đơn cho từng bữa ăn và xem danh sách "
        "nguyên liệu cần chuẩn bị tự động tổng hợp từ kế hoạch ăn uống.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Màn hình AI Coach Chatbot: Tích hợp khung trò chuyện trực tiếp với trợ lý AI, hỗ trợ giải đáp các thắc mắc "
        "về chế độ ăn kiêng, tập luyện và đưa ra lời khuyên dinh dưỡng hữu ích.", 
        list_bullet=True)

    # 2.4.2. Giao diện quản trị
    add_heading_styled(doc, "2.4.2. Giao diện quản trị", 2)
    
    add_custom_paragraph(doc, 
        "Phân hệ giao diện dành cho Quản trị viên (Admin Dashboard) cung cấp đầy đủ các công cụ kiểm soát "
        "và giám sát hoạt động hệ thống:")
        
    add_custom_paragraph(doc, 
        "Dashboard thống kê tổng quan: Hiển thị các chỉ số đo lường hiệu quả hoạt động bao gồm tổng số người dùng, "
        "tổng số lượt quét AI thành công, số lượng tài khoản hoạt động và các biểu đồ tăng trưởng số lượt quét theo thời gian.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Quản lý người dùng: Bảng danh sách chi tiết tài khoản hỗ trợ tìm kiếm theo tên/email. Cung cấp các nút hành động "
        "cho phép khóa tài khoản vi phạm, mở khóa tài khoản và phân quyền trực tiếp (nâng cấp thành ADMIN hoặc hạ cấp xuống USER).", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Quản lý Thư viện Thực phẩm (Food Library): Cho phép Admin quản lý danh mục thực phẩm dùng chung, hỗ trợ thêm món ăn mới "
        "kèm thông số dinh dưỡng chuẩn hóa, chỉnh sửa thông số món ăn hiện tại hoặc xóa khỏi thư viện.", 
        list_bullet=True)
        
    add_custom_paragraph(doc, 
        "Quản lý lịch sử quét: Bảng hiển thị toàn bộ lịch sử quét của người dùng trên toàn hệ thống kèm chức năng tìm kiếm "
        "và xem chi tiết ảnh quét cũng như kết quả JSON trả về.", 
        list_bullet=True)

    # Save document
    filename = "Chuong2_Noi_Dung_Thuc_Tap.docx"
    doc.save(filename)
    print(f"Chapter 2 report generated successfully: {filename}")

if __name__ == '__main__':
    main()
